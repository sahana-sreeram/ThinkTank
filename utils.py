from datetime import datetime
from typing import Dict, List, Any
import tempfile
import re, os, copy, shutil, subprocess, tempfile
import pandas as pd
import pypandoc

from docx import Document
from docx.shared import Inches, Pt


def _ensure_pandoc() -> None:
    """Lazily ensure pandoc is available (legacy DOCX path only). Avoids a network
    download at import time, which would break offline/local-first startup."""
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        pypandoc.download_pandoc()

import io
import json
from pathlib import Path

PROJECTS_FILE = Path("projects_db.json")
TEMPLATES_FILE = Path("scientist_templates.json")

def load_projects() -> Dict[str, Any]:
    if PROJECTS_FILE.exists():
        return json.loads(PROJECTS_FILE.read_text())
    return {}

def save_projects(data: Dict[str, Any]):
    PROJECTS_FILE.write_text(json.dumps(data, indent=2))

def load_templates() -> List[Dict[str, str]]:
    if TEMPLATES_FILE.exists():
        return json.loads(TEMPLATES_FILE.read_text())
    # initialize defaults if missing
    defaults = [
        {"title": "Immunologist", "expertise": "Immunopathology, antibody-antigen interactions",
         "goal": "Guide immune-targeting strategies", "role": "Analyse epitope selection and immune response"},
        {"title": "Machine Learning Expert", "expertise": "Deep learning, protein sequence modelling",
         "goal": "Develop predictive models for design", "role": "Build & chain ML models to rank candidates"},
        {"title": "Computational Biologist", "expertise": "Protein folding simulation, molecular dynamics",
         "goal": "Validate structural stability", "role": "Simulate docking & refine structures"}
    ]
    TEMPLATES_FILE.write_text(json.dumps(defaults, indent=2))
    return defaults

def save_templates(templates: List[Dict[str, str]]):
    TEMPLATES_FILE.write_text(json.dumps(templates, indent=2))

MERMAID_BLOCK = re.compile(r"```mermaid(.*?)```", re.S)

def _render_mermaid(code: str, out_png: str):
    """Render mermaid code to PNG with mermaid-cli (`mmdc`)."""
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False) as tmp:
        tmp.write(code)
        tmp_path = tmp.name
    try:
        subprocess.run(
            ["mmdc", "-i", tmp_path, "-o", out_png, "-b", "transparent"],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    finally:
        os.remove(tmp_path)

def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def clean_name(name: str, min_len: int = 3) -> str:
    # replace invalid chars with _
    clean = re.sub(r"[^A-Za-z0-9._-]", "_", name)
    # collapse consecutive underscores / dots / dashes
    clean = re.sub(r"[_\.-]{2,}", "_", clean)
    # strip leading / trailing non-alphanumerics
    clean = re.sub(r"^[^A-Za-z0-9]+|[^A-Za-z0-9]+$", "", clean)
    # fallback if too short
    if len(clean) < min_len:
        clean = (clean + "___")[:min_len]
    return clean[:512]

def indent(text: str, pad: int = 2) -> str:
    prefix = " " * pad
    return "\n".join(prefix + ln for ln in text.splitlines())

def _docx_bytes(project_name: str,
                project_desc: str,
                scientists: List[Dict[str, str]],
                md_text: List[str],
                table_font_size: int = 10) -> bytes:
    md_list = copy.deepcopy(md_text)        # never mutate caller’s list
    table_md = pd.DataFrame(scientists).to_markdown(index=False, tablefmt="pipe")
    md_list.insert(0, table_md)             # put table above any meeting notes
    # Add watermark at the end of list
    footer = (
        f'---\n\n')
    body_md   = "\n\n".join(md_list)
    header_md = (
        f"# {project_name}\n\n"
        f"---\n\n"
        f"## {project_desc}\n\n"
        f"**Exported on:** {now()}\n"
    )
    full_markdown = f"{header_md}\n\n{body_md}\n\n{footer}"

    images_dir = tempfile.mkdtemp()
    def _replace(match, counter=[0]):
        counter[0] += 1
        code = match.group(1).strip()
        img_path = os.path.join(images_dir, f"mermaid_{counter[0]}.png")
        _render_mermaid(code, img_path)
        return f"![diagram-{counter[0]}]({img_path})"

    full_md = MERMAID_BLOCK.sub(_replace, full_markdown)
    _ensure_pandoc()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    # Disable YAML metadata block parsing to avoid conflicts with --- separators
    pypandoc.convert_text(full_md, to="docx", format="markdown-yaml_metadata_block", outputfile=tmp_path)

    doc = Document(tmp_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(table_font_size)
    para = doc.add_paragraph()
    run = para.add_run()
    with open("assets/Logo_tau.png", "rb") as f: 
        pic_bytes = f.read()
    run.add_picture(io.BytesIO(pic_bytes), width=Inches(0.4))
    para.add_run("   Developed by TAU Group").font.size = Pt(14)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
        out_path = out.name
    doc.save(out_path)

    with open(out_path, "rb") as f:
        docx_bytes = f.read()

    os.remove(tmp_path)
    os.remove(out_path)
    shutil.rmtree(images_dir, ignore_errors=True)

    return docx_bytes

def export_meeting(project_name: str,
                   project_desc: str,
                   scientists: List[Dict[str, str]],
                   meeting: Dict[str, any],
                   md_text: List[str]) -> Dict[str, bytes]:
    """
    Returns {'docx': …, 'rtf': …, 'pdf': …}  - each value is file bytes.
    """
    return {
        "docx": _docx_bytes(project_name, project_desc, scientists, md_text)
    }

# ---------------------------------------------------------------------------
# Policy brief export (NEW — OWNER: Person 4). Pure python-docx, no pandoc.
# ---------------------------------------------------------------------------

def export_policy_brief(result) -> bytes:
    """Render a PolicyRunResult as a DOCX policy brief and return the bytes.

    Takes a models.PolicyRunResult. Self-contained (no pandoc/network) so it works
    offline. Used by the Streamlit download button.
    """
    doc = Document()
    req = result.request
    doc.add_heading("Policy Brief", level=0)
    doc.add_paragraph(req.question)
    meta = doc.add_paragraph()
    meta.add_run(f"Geography: {req.geography}   |   Generated: {now()}").italic = True

    rec = result.recommendation
    if rec:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(rec.summary)
        doc.add_paragraph(f"Confidence: {rec.confidence:.0%}")

        def _bullets(title, items):
            if not items:
                return
            doc.add_heading(title, level=2)
            for it in items:
                doc.add_paragraph(str(it), style="List Bullet")

        _bullets("Recommended Actions", rec.recommended_actions)
        _bullets("Benefits", rec.benefits)
        _bullets("Risks", rec.risks)
        _bullets("Equity Effects", rec.equity_effects)
        if rec.implementation_plan:
            doc.add_heading("Implementation Plan", level=2)
            for step in rec.implementation_plan.steps:
                doc.add_paragraph(
                    f"{step.phase} ({step.timeline or 'TBD'})", style="List Bullet"
                )
                for a in step.actions:
                    doc.add_paragraph(a, style="List Bullet 2")
        if rec.evidence_ids:
            doc.add_heading("Supporting Evidence", level=2)
            doc.add_paragraph(", ".join(sorted(set(rec.evidence_ids))))

    # Stakeholder positions
    if result.research:
        doc.add_heading("Stakeholder Views", level=1)
        for r in result.research:
            doc.add_heading(r.stakeholder, level=2)
            doc.add_paragraph(f"Position: {r.likely_position}")
            for f in r.findings:
                doc.add_paragraph(
                    f"{f.claim} [{', '.join(f.evidence_ids) or 'no citation'}]",
                    style="List Bullet",
                )

    # Forecast
    fc = result.forecast
    if fc:
        doc.add_heading("Forecast", level=1)
        if fc.mode == "qualitative":
            doc.add_paragraph("Qualitative outlook (no numeric model for this domain):")
            for line in fc.qualitative:
                doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph(f"Numeric scenarios ({fc.domain} domain):")
            for s in (fc.baseline, fc.conservative, fc.expected, fc.optimistic):
                if s is not None:
                    doc.add_paragraph(f"{s.name}: {s.inputs}", style="List Bullet")
        for a in fc.assumptions:
            doc.add_paragraph(f"Assumption: {a}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
